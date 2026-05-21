import logging
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from pypdf import PdfReader

from ..config import ENABLE_PDF_OCR_FALLBACK
from ..exceptions import ParseError


logger = logging.getLogger(__name__)
WORD_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
MIN_TEXT_LENGTH_FOR_DIRECT_AUDIT = 80


@dataclass
class ParseOutcome:
    text: str
    page_count: int | None = None
    readability: dict | None = None
    parser_used: str | None = None
    fallback_used: bool = False


def parse_contract_bytes(filename: str, content: bytes) -> tuple[str, int | None]:
    outcome = parse_contract_content(filename, content)
    return outcome.text, outcome.page_count


def parse_contract_content(filename: str, content: bytes) -> ParseOutcome:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(content)
    if suffix == ".docx":
        return _parse_docx(content)
    if suffix == ".doc":
        return _parse_doc(content)
    raise ParseError("UPLOAD_UNSUPPORTED_TYPE", "仅支持 PDF、DOC、DOCX 文件上传", 400, phase="upload.parse")


def normalize_text(text: str) -> str:
    normalized = text.replace("\ufeff", "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [line.strip() for line in normalized.split("\n")]
    compact = "\n".join(line for line in lines if line)
    return compact.strip()


def assess_text_readability(text: str) -> dict[str, float | bool | int]:
    sample = text.strip()
    length = len(sample)
    if not sample:
        return {
            "is_readable": False,
            "question_ratio": 1.0,
            "replacement_ratio": 0.0,
            "cjk_ratio": 0.0,
            "alnum_ratio": 0.0,
            "length": 0,
        }

    question_count = sample.count("?") + sample.count("锛?")
    replacement_count = sample.count("\ufffd")
    cjk_count = sum(1 for char in sample if "\u4e00" <= char <= "\u9fff")
    alnum_count = sum(1 for char in sample if char.isalnum())
    question_ratio = question_count / length
    replacement_ratio = replacement_count / length
    cjk_ratio = cjk_count / length
    alnum_ratio = alnum_count / length
    is_readable = not (
        replacement_ratio >= 0.02
        or question_ratio >= 0.35
        or (length >= 20 and alnum_ratio < 0.2 and cjk_ratio < 0.1)
    )
    return {
        "is_readable": is_readable,
        "question_ratio": round(question_ratio, 4),
        "replacement_ratio": round(replacement_ratio, 4),
        "cjk_ratio": round(cjk_ratio, 4),
        "alnum_ratio": round(alnum_ratio, 4),
        "length": length,
    }


def should_use_conservative_audit(text: str) -> bool:
    readability = assess_text_readability(text)
    return (not readability["is_readable"]) and readability["length"] < MIN_TEXT_LENGTH_FOR_DIRECT_AUDIT


def _parse_pdf(content: bytes) -> ParseOutcome:
    parser_attempts: list[tuple[str, callable]] = [
        ("pypdf", _extract_pdf_with_pypdf),
        ("raw-stream", _extract_pdf_with_raw_stream),
    ]
    if ENABLE_PDF_OCR_FALLBACK:
        parser_attempts.append(("ocr-fallback", _extract_pdf_with_ocr))
    errors: list[str] = []
    page_count: int | None = None
    best_text = ""
    best_parser: str | None = None

    for index, (parser_name, parser) in enumerate(parser_attempts):
        try:
            extracted_text, extracted_page_count = parser(content)
            if extracted_page_count is not None:
                page_count = extracted_page_count
            normalized = normalize_text(extracted_text)
            if not normalized:
                errors.append(f"{parser_name}: empty")
                continue
            readability = assess_text_readability(normalized)
            if len(normalized) > len(best_text):
                best_text = normalized
                best_parser = parser_name
            if readability["is_readable"] or len(normalized) >= 120:
                return ParseOutcome(
                    text=normalized,
                    page_count=page_count,
                    readability=readability,
                    parser_used=parser_name,
                    fallback_used=index > 0,
                )
            errors.append(f"{parser_name}: low-readability")
        except ParseError as exc:
            errors.append(f"{parser_name}: {exc.code}")
        except Exception as exc:
            logger.warning("pdf parser failed parser=%s error=%s", parser_name, exc)
            errors.append(f"{parser_name}: {exc.__class__.__name__}")

    if best_text:
        readability = assess_text_readability(best_text)
        return ParseOutcome(
            text=best_text,
            page_count=page_count,
            readability=readability,
            parser_used=best_parser,
            fallback_used=best_parser not in {None, "pypdf"},
        )

    raise ParseError(
        "PARSE_PDF_FAILED",
        "PDF 解析失败，请重新导出后上传，或改为直接粘贴文本",
        400,
        phase="upload.parse",
    )


def _parse_docx(content: bytes) -> ParseOutcome:
    parser_attempts: list[tuple[str, callable]] = [
        ("python-docx", _extract_docx_with_python_docx),
        ("docx-xml", _extract_docx_with_xml_fallback),
    ]
    errors: list[str] = []
    best_text = ""
    best_parser: str | None = None

    for index, (parser_name, parser) in enumerate(parser_attempts):
        try:
            extracted_text = normalize_text(parser(content))
            if not extracted_text:
                errors.append(f"{parser_name}: empty")
                continue
            readability = assess_text_readability(extracted_text)
            if len(extracted_text) > len(best_text):
                best_text = extracted_text
                best_parser = parser_name
            if readability["is_readable"] or len(extracted_text) >= 120:
                return ParseOutcome(
                    text=extracted_text,
                    page_count=None,
                    readability=readability,
                    parser_used=parser_name,
                    fallback_used=index > 0,
                )
            errors.append(f"{parser_name}: low-readability")
        except Exception as exc:
            logger.warning("docx parser failed parser=%s error=%s", parser_name, exc)
            errors.append(f"{parser_name}: {exc.__class__.__name__}")

    if best_text:
        readability = assess_text_readability(best_text)
        return ParseOutcome(
            text=best_text,
            page_count=None,
            readability=readability,
            parser_used=best_parser,
            fallback_used=best_parser not in {None, "python-docx"},
        )

    raise ParseError(
        "PARSE_DOCX_FAILED",
        "Word 解析失败，请重新导出后上传，或改为直接粘贴文本",
        400,
        phase="upload.parse",
    )


def _parse_doc(content: bytes) -> ParseOutcome:
    parser_attempts: list[tuple[str, callable]] = [
        ("soffice-convert", _extract_doc_via_soffice),
        ("antiword", _extract_doc_via_antiword),
        ("binary-text", _extract_doc_with_binary_fallback),
    ]
    errors: list[str] = []
    best_text = ""
    best_parser: str | None = None

    for index, (parser_name, parser) in enumerate(parser_attempts):
        try:
            extracted_text = normalize_text(parser(content))
            if not extracted_text:
                errors.append(f"{parser_name}: empty")
                continue
            readability = assess_text_readability(extracted_text)
            if len(extracted_text) > len(best_text):
                best_text = extracted_text
                best_parser = parser_name
            if readability["is_readable"] or len(extracted_text) >= 120:
                return ParseOutcome(
                    text=extracted_text,
                    page_count=None,
                    readability=readability,
                    parser_used=parser_name,
                    fallback_used=index > 0,
                )
            errors.append(f"{parser_name}: low-readability")
        except Exception as exc:
            logger.warning("doc parser failed parser=%s error=%s", parser_name, exc)
            errors.append(f"{parser_name}: {exc.__class__.__name__}")

    if best_text:
        readability = assess_text_readability(best_text)
        return ParseOutcome(
            text=best_text,
            page_count=None,
            readability=readability,
            parser_used=best_parser,
            fallback_used=True,
        )

    raise ParseError(
        "PARSE_DOC_FAILED",
        "DOC 解析失败，请优先转换为 DOCX 后重新上传，或改为直接粘贴文本",
        400,
        phase="upload.parse",
    )


def _extract_pdf_with_pypdf(content: bytes) -> tuple[str, int | None]:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages), len(reader.pages)


def _extract_pdf_with_raw_stream(content: bytes) -> tuple[str, int | None]:
    try:
        raw = content.decode("latin-1", errors="ignore")
    except Exception:
        return "", None
    tokens = re.findall(r"\(([^()]*)\)\s*Tj", raw)
    lines = [token.encode("latin-1", errors="ignore").decode("utf-8", errors="ignore") for token in tokens]
    return "\n".join(line for line in lines if line.strip()), None


def _extract_pdf_with_ocr(content: bytes) -> tuple[str, int | None]:
    tesseract_path = shutil.which("tesseract")
    pdftoppm_path = shutil.which("pdftoppm")
    if not tesseract_path or not pdftoppm_path:
        return "", None

    with tempfile.TemporaryDirectory() as temp_dir:
        pdf_path = Path(temp_dir) / "source.pdf"
        image_prefix = Path(temp_dir) / "page"
        pdf_path.write_bytes(content)

        convert_cmd = [pdftoppm_path, "-png", str(pdf_path), str(image_prefix)]
        convert_result = subprocess.run(convert_cmd, capture_output=True, text=True, timeout=60, check=False)
        if convert_result.returncode != 0:
            logger.warning("pdf ocr image conversion failed stderr=%s", convert_result.stderr.strip())
            return "", None

        texts: list[str] = []
        page_count = 0
        for image_path in sorted(Path(temp_dir).glob("page-*.png")):
            page_count += 1
            output_prefix = image_path.with_suffix("")
            ocr_cmd = [tesseract_path, str(image_path), str(output_prefix), "-l", "chi_sim+eng"]
            ocr_result = subprocess.run(ocr_cmd, capture_output=True, text=True, timeout=120, check=False)
            if ocr_result.returncode != 0:
                logger.warning("pdf ocr failed image=%s stderr=%s", image_path.name, ocr_result.stderr.strip())
                continue
            txt_path = output_prefix.with_suffix(".txt")
            if txt_path.exists():
                texts.append(txt_path.read_text(encoding="utf-8", errors="ignore"))
        return "\n".join(texts), page_count or None


def _extract_docx_with_python_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    segments = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                segments.append(" | ".join(cells))
    return "\n".join(segments)


def _extract_docx_with_xml_fallback(content: bytes) -> str:
    with zipfile.ZipFile(BytesIO(content)) as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml_bytes)
    paragraphs: list[str] = []
    for paragraph in root.findall(".//w:p", WORD_NAMESPACE):
        text_nodes = paragraph.findall(".//w:t", WORD_NAMESPACE)
        line = "".join(node.text or "" for node in text_nodes).strip()
        if line:
            paragraphs.append(line)
    return "\n".join(paragraphs)


def _extract_doc_via_soffice(content: bytes) -> str:
    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_path:
        return ""

    with tempfile.TemporaryDirectory() as temp_dir:
        source_path = Path(temp_dir) / "source.doc"
        source_path.write_bytes(content)
        convert_cmd = [
            soffice_path,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            temp_dir,
            str(source_path),
        ]
        result = subprocess.run(convert_cmd, capture_output=True, text=True, timeout=120, check=False)
        if result.returncode != 0:
            logger.warning("doc convert via soffice failed stderr=%s", result.stderr.strip())
            return ""
        converted_path = Path(temp_dir) / "source.docx"
        if not converted_path.exists():
            return ""
        return _extract_docx_with_python_docx(converted_path.read_bytes())


def _extract_doc_via_antiword(content: bytes) -> str:
    antiword_path = shutil.which("antiword")
    if not antiword_path:
        return ""
    with tempfile.TemporaryDirectory() as temp_dir:
        source_path = Path(temp_dir) / "source.doc"
        source_path.write_bytes(content)
        result = subprocess.run(
            [antiword_path, str(source_path)],
            capture_output=True,
            timeout=60,
            check=False,
        )
        if result.returncode != 0:
            logger.warning("doc convert via antiword failed stderr=%s", result.stderr.decode("utf-8", errors="ignore").strip())
            return ""
        return result.stdout.decode("utf-8", errors="ignore")


def _extract_doc_with_binary_fallback(content: bytes) -> str:
    decoded = content.decode("latin-1", errors="ignore")
    chunks = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：、“”‘’（）()\[\]《》\-_/#%&+,.:;!? ]{6,}", decoded)
    return "\n".join(chunks)
